"""Document service: Handle document upload, text extraction, and database updates."""
import os
import io
import logging
from typing import Optional
from fastapi import UploadFile
from sqlalchemy.orm import Session
from app.models.models import Document, DocumentStatus
from app.config import settings

logger = logging.getLogger(__name__)


async def save_and_extract_document(
    db: Session,
    upload_file: UploadFile,
    uploaded_by: Optional[str] = None,
    title: Optional[str] = None,
    language: str = "en"
) -> Document:
    """Save an uploaded document to disk, extract its text, and store in DB."""
    
    # 1. Provide an initial DB entry
    db_doc = Document(
        filename=upload_file.filename,
        title=title or upload_file.filename,
        file_path="",  # placeholder
        file_type="unknown",
        status=DocumentStatus.UPLOADED,
        language=language,
        uploaded_by=uploaded_by
    )
    db.add(db_doc)
    db.commit()
    db.refresh(db_doc)
    
    ext = os.path.splitext(upload_file.filename)[1].lower()
    file_type = "pdf" if ext == ".pdf" else "docx" if ext == ".docx" else "txt" if ext == ".txt" else "unknown"
    
    # 2. Save file
    file_path = os.path.join(settings.upload_dir, f"{db_doc.id}{ext}")
    content = await upload_file.read()
    
    with open(file_path, "wb") as f:
        f.write(content)
        
    db_doc.file_path = file_path
    db_doc.file_type = file_type
    db_doc.status = DocumentStatus.PROCESSING
    db.commit()
    
    try:
        # 3. Extract text
        extracted_text = ""
        if file_type == "pdf":
            import pypdf
            pdf_reader = pypdf.PdfReader(io.BytesIO(content))
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    extracted_text += text + "\n"
        elif file_type == "txt":
            extracted_text = content.decode("utf-8", errors="ignore")
        elif file_type == "docx":
            import docx
            doc = docx.Document(io.BytesIO(content))
            extracted_text = "\n".join([p.text for p in doc.paragraphs])
        else:
            raise ValueError(f"Unsupported file type: {ext}")
            
        # We don't store the full text in the DB to save space; 
        # it goes into ChromaDB.
        
        # 4. Integrate with RAG (vectorizing) -> done asynchronously later
        return db_doc, extracted_text
        
    except Exception as e:
        logger.error(f"Error processing document {db_doc.id}: {e}")
        db_doc.status = DocumentStatus.FAILED
        db_doc.summary = f"Error extracting text: {str(e)}"
        db.commit()
        raise e
