"""Export router: PDF, DOCX, JSON, Text export."""
import json
from io import BytesIO
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from app.database import get_db
from app.models.models import Meeting, Summary, ActionItem, User
from app.schemas import ExportRequest
from app.middleware.auth import get_current_user
from app.middleware.audit import log_action

router = APIRouter(prefix="/meetings", tags=["Export"])


@router.get("/{meeting_id}/export")
async def export_meeting(
    meeting_id: str,
    format: str = "json",
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    meeting = db.query(Meeting).filter(Meeting.id == meeting_id).first()
    if not meeting:
        raise HTTPException(status_code=404, detail="Meeting not found")

    summary = db.query(Summary).filter(Summary.meeting_id == meeting_id).first()
    actions = db.query(ActionItem).filter(ActionItem.meeting_id == meeting_id).all()

    log_action(db, current_user.id, "EXPORT", "meeting", meeting_id, {"format": format})

    if format == "json":
        return _export_json(meeting, summary, actions)
    elif format == "text":
        return _export_text(meeting, summary, actions)
    elif format == "pdf":
        return _export_pdf(meeting, summary, actions)
    elif format == "docx":
        return _export_docx(meeting, summary, actions)
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {format}")


def _export_json(meeting, summary, actions):
    data = {
        "meeting": {
            "title": meeting.title,
            "date": str(meeting.date),
            "venue": meeting.venue,
            "platform": meeting.platform,
            "attendees": meeting.attendees or [],
        },
        "summary": summary.content if summary else {},
        "action_items": [
            {
                "description": a.description,
                "assigned_to": a.assigned_to,
                "deadline": str(a.deadline) if a.deadline else None,
                "status": a.status.value,
                "priority": a.priority,
            }
            for a in actions
        ],
    }
    return data


def _export_text(meeting, summary, actions):
    lines = []
    lines.append("=" * 60)
    lines.append(f"MEETING SUMMARY REPORT")
    lines.append("=" * 60)
    lines.append(f"📌 Title: {meeting.title}")
    lines.append(f"📅 Date: {meeting.date}")
    lines.append(f"📍 Venue: {meeting.venue or meeting.platform}")
    lines.append(f"👥 Attendees: {', '.join(a.get('name', '') for a in (meeting.attendees or []))}")
    lines.append("")

    if summary and summary.content:
        content = summary.content

        if content.get("agenda_items"):
            lines.append("🗂 AGENDA ITEMS:")
            for i, item in enumerate(content["agenda_items"], 1):
                lines.append(f"  {i}. {item}")
            lines.append("")

        if content.get("key_points"):
            lines.append("📝 KEY DISCUSSION POINTS:")
            for point in content["key_points"]:
                lines.append(f"  • {point}")
            lines.append("")

        if content.get("decisions"):
            lines.append("✅ DECISIONS TAKEN:")
            for dec in content["decisions"]:
                lines.append(f"  • {dec}")
            lines.append("")

    if actions:
        lines.append("📌 ACTION ITEMS:")
        lines.append(f"  {'Action':<40} {'Assigned To':<20} {'Deadline':<15} {'Status':<10}")
        lines.append(f"  {'-'*40} {'-'*20} {'-'*15} {'-'*10}")
        for a in actions:
            lines.append(
                f"  {a.description[:40]:<40} {a.assigned_to:<20} "
                f"{str(a.deadline)[:10] if a.deadline else 'N/A':<15} {a.status.value:<10}"
            )
        lines.append("")

    if summary and summary.content:
        if summary.content.get("flagged_items"):
            lines.append("⚠️ FLAGGED ITEMS:")
            for item in summary.content["flagged_items"]:
                lines.append(f"  • {item}")
            lines.append("")

    lines.append("=" * 60)
    text = "\n".join(lines)

    return StreamingResponse(
        BytesIO(text.encode("utf-8")),
        media_type="text/plain",
        headers={"Content-Disposition": f'attachment; filename="meeting_{meeting.id[:8]}.txt"'}
    )


def _export_pdf(meeting, summary, actions):
    """Generate a PDF export using fpdf2."""
    try:
        from fpdf import FPDF
    except ImportError:
        raise HTTPException(status_code=500, detail="PDF library not installed")

    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Title
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Meeting Summary Report", ln=True, align="C")
    pdf.ln(5)

    # Meeting info
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 7, f"Title: {meeting.title}", ln=True)
    pdf.cell(0, 7, f"Date: {meeting.date}", ln=True)
    pdf.cell(0, 7, f"Venue: {meeting.venue or meeting.platform}", ln=True)
    pdf.ln(5)

    # Summary content
    if summary and summary.content:
        content = summary.content

        if content.get("key_points"):
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, "Key Discussion Points", ln=True)
            pdf.set_font("Helvetica", "", 10)
            for point in content["key_points"]:
                pdf.multi_cell(0, 6, f"  - {point}")
            pdf.ln(3)

        if content.get("decisions"):
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 8, "Decisions Taken", ln=True)
            pdf.set_font("Helvetica", "", 10)
            for dec in content["decisions"]:
                pdf.multi_cell(0, 6, f"  - {dec}")
            pdf.ln(3)

    # Action items
    if actions:
        pdf.set_font("Helvetica", "B", 12)
        pdf.cell(0, 8, "Action Items", ln=True)
        pdf.set_font("Helvetica", "", 9)
        for a in actions:
            pdf.multi_cell(0, 5,
                f"  [{a.status.value.upper()}] {a.description} | "
                f"Assigned: {a.assigned_to} | "
                f"Deadline: {str(a.deadline)[:10] if a.deadline else 'N/A'}"
            )

    buf = BytesIO(pdf.output())
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="meeting_{meeting.id[:8]}.pdf"'}
    )


def _export_docx(meeting, summary, actions):
    """Generate a DOCX export."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        raise HTTPException(status_code=500, detail="DOCX library not installed")

    doc = Document()
    doc.add_heading("Meeting Summary Report", 0)

    doc.add_paragraph(f"Title: {meeting.title}")
    doc.add_paragraph(f"Date: {meeting.date}")
    doc.add_paragraph(f"Venue: {meeting.venue or meeting.platform}")

    if summary and summary.content:
        content = summary.content
        if content.get("key_points"):
            doc.add_heading("Key Discussion Points", level=1)
            for point in content["key_points"]:
                doc.add_paragraph(point, style="List Bullet")

        if content.get("decisions"):
            doc.add_heading("Decisions Taken", level=1)
            for dec in content["decisions"]:
                doc.add_paragraph(dec, style="List Bullet")

    if actions:
        doc.add_heading("Action Items", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Action"
        hdr[1].text = "Assigned To"
        hdr[2].text = "Deadline"
        hdr[3].text = "Status"
        for a in actions:
            row = table.add_row().cells
            row[0].text = a.description
            row[1].text = a.assigned_to
            row[2].text = str(a.deadline)[:10] if a.deadline else "N/A"
            row[3].text = a.status.value

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="meeting_{meeting.id[:8]}.docx"'}
    )
